"""Text extraction from blob bytes, dispatched by file extension.

Deliberately dependency-light: pypdf + python-docx + stdlib covers the vast
majority of real blob content. Heavier parsers (Tika, OCR) belong behind an
optional compose profile, not in the core image.
"""
from __future__ import annotations

import io
import logging

from pypdf import PdfReader
from docx import Document

log = logging.getLogger("bloblens.extract")

PLAIN_TEXT_EXTENSIONS = {
    "txt", "md", "csv", "tsv", "json", "log", "yml", "yaml",
    "xml", "html", "htm", "ini", "conf", "toml", "env", "sql", "py",
    "js", "ts", "sh", "ps1", "cs", "java", "go", "rb", "php",
}


def extension_of(blob_name: str) -> str:
    if "." not in blob_name.rsplit("/", 1)[-1]:
        return ""
    return blob_name.rsplit(".", 1)[-1].lower()


def is_extractable(blob_name: str) -> bool:
    ext = extension_of(blob_name)
    return ext in PLAIN_TEXT_EXTENSIONS or ext in ("pdf", "docx")


def extract_text(blob_name: str, data: bytes, max_chars: int) -> str:
    """Return extracted text truncated to max_chars. Never raises."""
    ext = extension_of(blob_name)
    try:
        if ext == "pdf":
            text = _from_pdf(data, max_chars)
        elif ext == "docx":
            text = _from_docx(data)
        elif ext in PLAIN_TEXT_EXTENSIONS:
            text = data.decode("utf-8", errors="ignore")
        else:
            return ""
        return text[:max_chars]
    except Exception as exc:  # noqa: BLE001 - indexing must survive bad files
        log.warning("extraction failed for %s: %s", blob_name, exc)
        return ""


def _from_pdf(data: bytes, max_chars: int) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    total = 0
    for page in reader.pages:
        chunk = page.extract_text() or ""
        parts.append(chunk)
        total += len(chunk)
        if total >= max_chars:
            break
    return "\n".join(parts)


def _from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)
